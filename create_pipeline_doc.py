"""Generate Pipeline V2.1 Algorithm Reference DOCX."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

# --- Helper functions ---
def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.5 * level)
    return p

def add_formula(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

def add_output_box(title, items):
    """Add a shaded output expectation box."""
    add_para(title, bold=True)
    for item in items:
        add_bullet(item)

# ============================================================
# DOCUMENT CONTENT
# ============================================================

# Title Page
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\n')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Pipeline V2.1')
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Temporal Contiguity Analysis\nfor Instructional Videos')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\nComplete Algorithm Reference\n& Expected Stage Outputs')
run.font.size = Pt(16)
run.italic = True
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\nVersion 2.1 (GPU)\nMarch 2026')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ============================================================
# OVERVIEW
# ============================================================
add_heading('1. Pipeline Overview', 1)

add_para('The pipeline measures Temporal Contiguity \u2014 how well visual content and narration are time-aligned in instructional videos. It answers the question: "When the narrator mentions a concept, how close in time is the corresponding visual?"')

add_para('Core Metric:', bold=True)
add_formula('delta_t = t_narr \u2212 t_vis')
add_para('Where t_narr = when the narrator says it, t_vis = when the viewer first sees it.')

add_para('A small |delta_t| means good temporal contiguity (visual and audio are synchronized). A large |delta_t| means the learner must hold information in working memory, increasing cognitive load.')

add_heading('1.1 Stage Sequence', 2)
add_table(
    ['Stage', 'Name', 'Purpose'],
    [
        ['1', 'ASR Transcription', 'Extract word-level transcript with timestamps from audio'],
        ['2', 'Scene Detection', 'Find visual scene boundaries using 3-signal fusion'],
        ['3', 'Visual Concepts', 'Label each scene with OCR text + content type'],
        ['4a', 'Text Embeddings', 'Pre-compute vector representations of all transcript segments'],
        ['4b', 'Narration Alignment', 'Match each scene to its best narration segment (3-track cascade)'],
        ['5', 'Keyword Grounding', 'Extract keywords, find where each appears visually (4-step cascade)'],
        ['6', 'Pedagogical Importance', 'Rate each segment\u2019s educational importance (1\u20135)'],
        ['7', 'Scoring & Aggregation', 'Compute S_temporal scores, zones, and overall grade'],
        ['\u2014', 'Dashboard', 'Generate HTML report with per-scene breakdown'],
    ],
    col_widths=[0.6, 1.8, 4.0]
)

add_heading('1.2 Models Used', 2)
add_table(
    ['Model', 'Purpose', 'Dimensions'],
    [
        ['faster-whisper medium', 'Speech-to-text with word timestamps', '\u2014'],
        ['facebook/dinov2-base', 'Visual frame embeddings for scene boundaries', '768-dim'],
        ['BAAI/bge-large-en-v1.5', 'Text embeddings for semantic matching', '1024-dim'],
        ['ViT-B-16-SigLIP (webli)', 'Vision-text alignment + frame classification', '512-dim'],
        ['EasyOCR', 'Optical character recognition on frames', '\u2014'],
        ['spaCy en_core_web_sm', 'Keyword extraction (noun chunks, NER)', '\u2014'],
    ],
    col_widths=[2.2, 3.0, 1.2]
)

doc.add_page_break()

# ============================================================
# STAGE 1: ASR
# ============================================================
add_heading('2. Stage 1: ASR Transcription', 1)

add_para('Purpose: Extract a word-level transcript with precise timestamps from the video\u2019s audio track.', bold=True)

add_heading('2.1 Algorithm', 2)
add_bullet('Extract audio from video as 16 kHz mono WAV using ffmpeg')
add_bullet('Try WhisperX (medium model) with forced alignment for per-word timestamps')
add_bullet('If WhisperX fails (e.g., PyTorch version mismatch) \u2192 fall back to faster-whisper with word_timestamps=True')
add_bullet('GPU auto-detection: float16 on CUDA, int8 on CPU')
add_bullet('Re-segment transcript into proper sentences using pySBD (Python Sentence Boundary Detection)')
add_bullet('Merge short segments (< 4 words) into predecessor to avoid micro-segments')
add_bullet('Run timestamp validation: flag words with duration outside [0.05s, 2.0s], gaps > 3s, or non-monotonic timestamps')

add_heading('2.2 Key Parameters', 2)
add_table(
    ['Parameter', 'Value', 'Meaning'],
    [
        ['WHISPER_MODEL', 'medium', 'Whisper model size (trade-off: speed vs accuracy)'],
        ['WHISPER_COMPUTE_TYPE', 'float16', 'GPU precision (int8 on CPU)'],
        ['WHISPER_BATCH_SIZE', '16', 'Parallel audio chunks on GPU'],
        ['Min merge words', '4', 'Segments with < 4 words merge into predecessor'],
    ],
    col_widths=[2.0, 1.2, 3.2]
)

add_heading('2.3 Expected Output', 2)
add_para('File: transcript_words.csv', bold=True)
add_table(
    ['word_id', 'word', 'start_time', 'end_time', 'timestamp_reliable'],
    [
        ['0', 'Imagine', '3.680', '4.020', 'True'],
        ['1', "you're", '4.020', '4.200', 'True'],
        ['2', 'standing', '4.200', '4.560', 'True'],
        ['3', 'in', '4.560', '4.640', 'True'],
        ['...', '...', '...', '...', '...'],
    ],
    col_widths=[0.8, 1.2, 1.2, 1.2, 1.5]
)
add_para('Expect: 400\u2013800 words for a 3\u20134 minute video. Each word has a precise start/end time. timestamp_reliable=False flags words with suspicious timing.', italic=True)

add_para('\nFile: transcript_segments_improved.csv', bold=True)
add_table(
    ['segment_id', 'text', 'start_time', 'end_time', 'n_words'],
    [
        ['0', 'Imagine you\'re standing in a foggy landscape...', '3.68', '10.46', '18'],
        ['1', 'This idea captures the essence of local search...', '10.46', '14.24', '11'],
        ['...', '...', '...', '...', '...'],
    ],
    col_widths=[0.8, 3.0, 0.9, 0.9, 0.8]
)
add_para('Expect: 25\u201360 sentences for a 3\u20134 minute video. Proper sentence boundaries, not arbitrary Whisper chunks.', italic=True)

doc.add_page_break()

# ============================================================
# STAGE 2: SCENE DETECTION
# ============================================================
add_heading('3. Stage 2: Scene Detection (Multi-Signal Fusion)', 1)

add_para('Purpose: Detect visual scene boundaries by fusing three complementary signals. Each scene represents a visually coherent segment of the video.', bold=True)

add_heading('3.1 Frame Extraction', 2)
add_bullet('Sample frames at 0.5s intervals (2 fps) using ffmpeg')
add_bullet('A 3.5-minute video produces ~433 frames')
add_bullet('Each frame saved as JPEG with timestamp in filename')

add_heading('3.2 Signal A: PySceneDetect (weight = 0.35)', 2)
add_bullet('Uses AdaptiveDetector with window_width=5')
add_bullet('Detects hard cuts and transitions')
add_bullet('Produces binary signal: 1.0 at detected boundaries, 0.0 elsewhere')
add_bullet('Good at: sharp visual transitions (slide changes, camera cuts)')
add_bullet('Bad at: gradual animations, slight content changes within same visual style')

add_heading('3.3 Signal B: DINOv2 Embedding Distance (weight = 0.45)', 2)
add_bullet('Model: facebook/dinov2-base (768-dim CLS token)')
add_bullet('Compute embedding for each frame, then cosine distance between consecutive frames')
add_formula('distance(i) = 1 \u2212 cos_sim(frame[i], frame[i+1])')
add_bullet('Normalize to [0, 1] by dividing by max distance')
add_bullet('Good at: detecting semantic visual changes (different diagram, new concept)')
add_bullet('Bad at: may miss text-only changes on same background')

add_heading('3.4 Signal C: OCR Jaccard Distance (weight = 0.20)', 2)
add_bullet('Run OCR every ~1.5s (every 3rd frame at 0.5s interval)')
add_bullet('Extract word set from each sampled frame')
add_bullet('Compute Jaccard distance between consecutive samples:')
add_formula('J(A, B) = 1 \u2212 |A \u2229 B| / |A \u222a B|')
add_bullet('Only flag if distance > 0.5 (OCR_JACCARD_THRESHOLD)')
add_bullet('Spread signal across skipped frames, normalize to [0, 1]')
add_bullet('Good at: detecting text content changes (new slide with different text)')
add_bullet('Bad at: frames with no text; two different slides with similar words')

add_heading('3.5 Signal Fusion', 2)
add_para('If OCR signal has non-zero values (text was detected):')
add_formula('combined = 0.35 \u00d7 A + 0.45 \u00d7 B + 0.20 \u00d7 C')
add_para('If no OCR text detected (all-zero signal C):')
add_formula('combined = 0.4375 \u00d7 A + 0.5625 \u00d7 B')

add_heading('3.6 Adaptive Threshold (Otsu-like)', 2)
add_bullet('Try 50 candidate thresholds from min to max of combined signal')
add_bullet('For each: compute between-class variance w0 \u00d7 w1 \u00d7 (\u03bc0 \u2212 \u03bc1)\u00b2')
add_bullet('Pick threshold with maximum variance')
add_bullet('Fallback: mean + 1.5 \u00d7 std')
add_bullet('Guard: if < 3 boundaries detected and video > 60s, lower threshold by 0.5 \u00d7 std')

add_heading('3.7 Boundary Detection + Scene Construction', 2)
add_bullet('Find frames where combined signal \u2265 threshold')
add_bullet('Non-maximum suppression: group adjacent frames (distance \u2264 2), keep peak')
add_bullet('Build scenes between consecutive boundaries')
add_bullet('Keyframe selection: DINOv2 centroid of inner 80% of frames (avoids first/last 10%)')
add_bullet('Run OCR on keyframe to capture scene\u2019s text content')
add_bullet('Smart merge: if scene < 2.0s AND DINOv2 similarity with predecessor > 0.85, merge')

add_heading('3.8 Expected Output', 2)
add_para('File: scenes.csv', bold=True)
add_table(
    ['scene_id', 't_start', 't_end', 'duration', 't_keyframe', 'ocr_words', 'n_ocr_words'],
    [
        ['0', '0.0', '3.5', '3.5', '1.5', '', '0'],
        ['1', '3.5', '28.0', '24.5', '15.0', 'local search optimization', '3'],
        ['2', '28.0', '55.0', '27.0', '42.0', 'state current solution', '3'],
        ['...', '...', '...', '...', '...', '...', '...'],
    ],
    col_widths=[0.7, 0.7, 0.7, 0.8, 0.9, 1.8, 0.8]
)
add_para('Expect: 10\u201325 scenes for a 3\u20134 min video. Ideal scene duration: 5\u201315s. Scenes > 30s suggest under-segmentation. t_keyframe should be near the middle of the scene.', italic=True)

add_para('\nFile: ocr_per_frame.csv', bold=True)
add_table(
    ['frame_time', 'words', 'n_words', 'mean_confidence'],
    [
        ['0.0', '', '0', '0.0'],
        ['1.5', 'al counts', '2', '0.91'],
        ['3.0', 'local search optimization algorithms', '4', '0.88'],
        ['...', '...', '...', '...'],
    ],
    col_widths=[1.0, 3.0, 0.8, 1.4]
)
add_para('Expect: One row per frame (433 rows for 3.5min video). Many frames may have 0 OCR words (animations, diagrams without text). Text slides should show captured words.', italic=True)

doc.add_page_break()

# ============================================================
# STAGE 3: VISUAL CONCEPTS
# ============================================================
add_heading('4. Stage 3: Visual Concept Extraction', 1)

add_para('Purpose: For each scene, determine what visual content is present (text, diagrams, etc.) and classify whether it\u2019s educational content or non-content (title screens, transitions).', bold=True)

add_heading('4.1 Algorithm', 2)
add_bullet('For each scene, run OCR on the keyframe to extract visible text')
add_bullet('Optionally run VLM (Gemini/Ollama) for description \u2014 currently set to "skip"')
add_bullet('Build concept_text = OCR text + VLM description')
add_bullet('Classify frame type using SigLIP zero-shot classification:')

add_para('Content labels: "content slide", "diagram", "code editor", "whiteboard", "animation frame", "demonstration"', italic=True)
add_para('Non-content labels: "title slide", "logo screen", "blank screen", "loading screen", "transition effect", "section divider", "talking head with no visual aids"', italic=True)

add_bullet('Compute sigmoid similarity between keyframe and all 13 labels')
add_bullet('Best label = argmax similarity; is_content = True if best is in content set')
add_bullet('Low-confidence override: if confidence < 0.55, default to content (conservative)')
add_bullet('OCR override: if any OCR text detected, force is_content = True')

add_heading('4.2 Expected Output', 2)
add_para('File: scene_concepts.csv', bold=True)
add_table(
    ['scene_id', 'ocr_text', 'concept_text', 'is_content', 'frame_type', 'frame_type_confidence'],
    [
        ['0', '', '', 'True', 'loading screen (low_conf)', '0.517'],
        ['1', 'Local Search and Optimization Algorithms', 'Local Search and...', 'True', 'content slide (low_conf)', '0.522'],
        ['2', 'Simplest Local Search Methods', 'Simplest Local...', 'True', 'content slide (low_conf)', '0.516'],
        ['...', '...', '...', '...', '...', '...'],
    ],
    col_widths=[0.6, 1.8, 1.5, 0.7, 1.4, 0.8]
)
add_para('Expect: One row per scene. Scenes with text slides should have rich ocr_text. Most scenes should be is_content=True. frame_type_confidence near 0.50 indicates SigLIP cannot confidently classify (all currently show "low_conf").', italic=True)

doc.add_page_break()

# ============================================================
# STAGE 4a: TEXT EMBEDDINGS
# ============================================================
add_heading('5. Stage 4a: Text Unit Embeddings', 1)

add_para('Purpose: Pre-compute dense vector representations of all transcript segments for use in alignment (Stages 4b and 5).', bold=True)

add_heading('5.1 Algorithm', 2)
add_bullet('Load transcript_segments_improved.csv')
add_bullet('Compute midpoint for each segment: t_mid = (start_time + end_time) / 2')
add_bullet('Embed each segment\u2019s text using BAAI/bge-large-en-v1.5 (1024-dim)')
add_bullet('Prefix each text with: "Represent this sentence: " (required by BGE)')
add_bullet('L2-normalize all embeddings')

add_heading('5.2 Expected Output', 2)
add_para('File: segment_meta.csv', bold=True)
add_table(
    ['segment_id', 'text', 'start_time', 'end_time', 't_mid'],
    [
        ['0', 'Imagine you\'re standing in a foggy...', '3.68', '10.46', '7.07'],
        ['1', 'This idea captures the essence...', '10.46', '14.24', '12.35'],
        ['...', '...', '...', '...', '...'],
    ],
    col_widths=[0.7, 3.0, 0.8, 0.8, 0.8]
)
add_para('File: segment_embeddings.npy \u2014 NumPy array of shape (N_segments, 1024). Each row is a normalized embedding vector.', italic=True)

doc.add_page_break()

# ============================================================
# STAGE 4b: ALIGNMENT
# ============================================================
add_heading('6. Stage 4b: Narration Alignment (3-Track Cascade)', 1)

add_para('Purpose: For each scene, find the best-matching narration segment. This is the CORE of temporal contiguity measurement \u2014 it determines delta_t.', bold=True)

add_heading('6.1 Cascade Logic', 2)
add_para('For each content scene, try tracks in priority order:')
add_bullet('Track A (OCR Word Match) \u2192 if matched, done')
add_bullet('Track B (SigLIP Vision-Text) \u2192 if matched, done')
add_bullet('Track C (Semantic Text Similarity) \u2192 if matched, done')
add_bullet('Else \u2192 no match (delta_t = None)')

add_heading('6.2 Track A: Exact OCR Word Matching', 2)
add_para('Confidence: HIGHEST (\u03b1 = 1.0)', bold=True)
add_bullet('Extract OCR words from scene (filter: length \u2265 3, not stop words)')
add_bullet('Search transcript words within \u00b130s of t_keyframe')
add_bullet('For each OCR word, find exact normalized match in transcript')
add_bullet('Record time offset: t_word \u2212 t_keyframe')
add_bullet('Context validation: embed concept_text, check cosine sim \u2265 0.25 with closest segment')
add_bullet('t_narr = t_keyframe + median(all word offsets)')
add_bullet('\u03b1 = 1.0 (maximum confidence for exact text match)')

add_heading('6.3 Track B: SigLIP Vision-to-Text', 2)
add_para('Confidence: MEDIUM (\u03b1 \u2208 [0.5, 1.0])', bold=True)
add_bullet('Embed keyframe with SigLIP vision encoder \u2192 512-dim vector')
add_bullet('Embed all transcript segments with SigLIP text encoder')
add_bullet('Compute sigmoid similarities (not softmax)')
add_bullet('Apply temporal Gaussian decay centered on t_keyframe:')
add_formula('weight(t) = exp(\u22120.5 \u00d7 ((t_mid \u2212 t_keyframe) / 15.0)\u00b2)')
add_bullet('Best segment = argmax(raw_sim \u00d7 temporal_weight)')
add_bullet('Minimum similarity threshold: 0.20')
add_bullet('\u03b1 = 0.5 + 0.5 \u00d7 (sim \u2212 0.20) / 0.40, clamped to [0.5, 1.0]')

add_heading('6.4 Track C: Semantic Text Similarity', 2)
add_para('Confidence: LOWEST (\u03b1 \u2208 [0.0, 1.0])', bold=True)
add_bullet('Embed concept_text with BGE-large \u2192 1024-dim')
add_bullet('Cosine similarity with all pre-computed segment embeddings')
add_bullet('Apply same temporal Gaussian decay (\u03c3 = 15.0s)')
add_bullet('Best segment = argmax(raw_sim \u00d7 temporal_weight)')
add_bullet('Minimum similarity: 0.20')
add_bullet('\u03b1 = (sim \u2212 0.30) / 0.50, clamped to [0.0, 1.0]')

add_heading('6.5 delta_t Computation', 2)
add_para('CRITICAL: After any track returns a result, the wrapper recomputes:', bold=True)
add_formula('t_vis = scene["t_start"]  (scene start time)')
add_formula('delta_t = t_narr \u2212 t_vis')
add_para('Note: All tracks internally match relative to t_keyframe, but the final delta_t is computed relative to t_start. For long scenes (e.g., 50+ seconds), t_start can be very far from t_keyframe, creating artificially large delta_t values. This is a known issue.', italic=True)

add_heading('6.6 Expected Output', 2)
add_para('File: alignment_events.csv', bold=True)
add_table(
    ['scene_id', 't_vis', 't_keyframe', 'match_track', 't_narr', 'delta_t', 'alpha', 'n_word_matches'],
    [
        ['0', '0.0', '1.5', 'B', '7.07', '7.07', '0.877', '0'],
        ['1', '3.5', '22.5', 'A', '20.83', '17.33', '1.0', '4'],
        ['3', '84.0', '84.5', 'B', '83.22', '\u22120.78', '0.871', '0'],
        ['4', '86.0', '88.0', 'B', '86.31', '0.31', '0.901', '0'],
        ['...', '...', '...', '...', '...', '...', '...', '...'],
    ],
    col_widths=[0.6, 0.6, 0.8, 0.8, 0.7, 0.7, 0.7, 0.8]
)
add_para('Expect: One row per scene. Track A for scenes with OCR text, Track B/C for others. delta_t should ideally be within \u00b13s. |delta_t| > 10s suggests a matching or scene detection problem. alpha = 1.0 for Track A, 0.5\u20131.0 for Track B, 0.0\u20131.0 for Track C.', italic=True)

doc.add_page_break()

# ============================================================
# STAGE 5: KEYWORDS
# ============================================================
add_heading('7. Stage 5: Keyword Extraction & Visual Grounding', 1)

add_para('Purpose: Extract keywords from narration, classify their visual groundability, then find where each keyword appears visually. Provides fine-grained keyword-level temporal contiguity.', bold=True)

add_heading('7.1 Keyword Extraction', 2)
add_bullet('spaCy (en_core_web_sm): extract noun chunks, named entities, standalone nouns')
add_bullet('KeyBERT: extract top-5 keyphrases using BGE embeddings (n-gram 1\u20133, MMR diversity 0.5)')
add_bullet('Merge and deduplicate: if single word is substring of multi-word phrase, keep only phrase')
add_bullet('Filter: length \u2265 3, not in stop words')

add_heading('7.2 Groundability Classification', 2)
add_table(
    ['Condition', 'Groundability', 'Example'],
    [
        ['In VISUAL_NOUNS set', 'HIGH', 'diagram, chart, code, formula, graph'],
        ['Concreteness score \u2265 4.0', 'HIGH', 'mountain, keyboard, screen'],
        ['Concreteness score \u2265 2.5', 'MEDIUM', 'algorithm, process'],
        ['In ABSTRACT_WORDS set', 'at most MEDIUM', 'understanding, concept, method'],
        ['Concreteness < 2.5', 'LOW \u2192 SKIP', 'importance, relationship'],
    ],
    col_widths=[2.0, 1.2, 3.2]
)

add_heading('7.3 4-Step Grounding Cascade', 2)

add_para('Step 1: OCR Fuzzy Search (confidence: HIGH)', bold=True)
add_bullet('Window: \u00b160s around t_narr')
add_bullet('Fuzzy match: Levenshtein ratio > 0.80')
add_bullet('Multi-word: ALL component words must match')
add_bullet('Pick closest frame to t_narr with match')

add_para('Step 2: GroundingDINO (confidence: MEDIUM) \u2014 disabled by default', bold=True)

add_para('Step 3: SigLIP Contextual (confidence: LOW)', bold=True)
add_bullet('Window: \u00b130s around t_narr')
add_bullet('Text query: "A visual showing {keyword}"')
add_bullet('Sigmoid similarity with temporal Gaussian decay (\u03c3 = 15s)')
add_bullet('Threshold: 0.15 (very lenient)')

add_para('Step 4: VLM Check (confidence: LOW) \u2014 only for HIGH groundability, currently skipped', bold=True)

add_heading('7.4 Expected Output', 2)
add_para('File: keyword_alignment.csv', bold=True)
add_table(
    ['keyword_text', 'segment_id', 't_narr', 't_vis', 'delta_t', 'method', 'confidence', 'groundability'],
    [
        ['local search', '1', '10.5', '3.0', '\u22127.5', 'ocr', 'HIGH', 'MEDIUM'],
        ['optimization', '1', '11.2', '3.0', '\u22128.2', 'ocr', 'HIGH', 'MEDIUM'],
        ['hill climbing', '5', '65.3', '63.0', '\u22122.3', 'siglip_contextual', 'LOW', 'MEDIUM'],
        ['understanding', '3', '30.0', 'None', 'None', 'skipped_low', 'NONE', 'LOW'],
        ['...', '...', '...', '...', '...', '...', '...', '...'],
    ],
    col_widths=[1.2, 0.7, 0.6, 0.6, 0.6, 1.2, 0.8, 0.9]
)
add_para('Expect: 100\u2013400 keywords for a 3\u20134 min video. ~10\u201320% grounded via OCR (highest quality), ~60\u201380% via SigLIP contextual (lower quality), ~10\u201320% skipped (LOW groundability or not found).', italic=True)

doc.add_page_break()

# ============================================================
# STAGE 6: IMPORTANCE
# ============================================================
add_heading('8. Stage 6: Pedagogical Importance Rating', 1)

add_para('Purpose: Rate each transcript segment\u2019s educational importance (1\u20135) to weight the final score. Critical content with poor timing matters more than filler with poor timing.', bold=True)

add_heading('8.1 3-Tier Backend', 2)
add_bullet('Tier 1: Gemini API (if API key available) \u2014 best quality, optional double-run for reliability')
add_bullet('Tier 2: Local LLM via Ollama (if running) \u2014 good quality, slower')
add_bullet('Tier 3: Heuristic (always available) \u2014 keyword density + speech rate + word count')

add_heading('8.2 Importance Scale', 2)
add_table(
    ['Rating', 'Level', 'Description', 'Weight'],
    [
        ['1', 'Low', 'Intro/outro, greetings, transitions, filler', '0.3'],
        ['2', 'Below Average', 'Recap of known material, tangential examples', '0.6'],
        ['3', 'Average', 'Supporting explanation, context building', '1.0'],
        ['4', 'Above Average', 'Key concept introduction, important examples', '1.5'],
        ['5', 'Critical', 'Core derivation, formula, definition, step-by-step', '2.0'],
    ],
    col_widths=[0.6, 1.2, 3.0, 0.6]
)

add_heading('8.3 Heuristic Formula (Tier 3)', 2)
add_formula('score = 0.4 \u00d7 kw_density + 0.3 \u00d7 speech_rate + 0.3 \u00d7 word_count_norm')
add_bullet('Map score to 1\u20135 via 20th/40th/60th/80th percentile bins')

add_heading('8.4 Expected Output', 2)
add_para('File: pedagogical_importance.csv', bold=True)
add_table(
    ['segment_id', 'importance', 'reason', 'backend', 'is_reliable'],
    [
        ['0', '3', 'heuristic', 'heuristic', 'False'],
        ['1', '4', 'heuristic', 'heuristic', 'False'],
        ['5', '5', 'heuristic', 'heuristic', 'False'],
        ['...', '...', '...', '...', '...'],
    ],
    col_widths=[0.8, 0.8, 1.5, 1.2, 0.8]
)
add_para('Expect: One row per segment. With heuristic backend, is_reliable=False (no LLM validation). Distribution should spread across 1\u20135 levels.', italic=True)

doc.add_page_break()

# ============================================================
# STAGE 7: SCORING
# ============================================================
add_heading('9. Stage 7: Scoring & Aggregation', 1)

add_para('Purpose: Convert delta_t values into temporal contiguity scores (0\u2013100) and classify into cognitive zones. Compute overall video grade.', bold=True)

add_heading('9.1 Temporal Scoring Formula', 2)
add_formula('S_temporal = 100 \u00d7 exp(\u22120.5 \u00d7 (|delta_t| / \u03c4)\u00b2)     \u03c4 = 2.5s')

add_table(
    ['|delta_t|', 'S_temporal', 'Zone', 'Interpretation'],
    [
        ['0.0s', '100.0', 'Optimal', 'Perfect synchronization'],
        ['0.5s', '98.0', 'Optimal', 'Nearly perfect'],
        ['1.0s', '92.3', 'Optimal (\u22641s)', 'Excellent alignment'],
        ['2.0s', '72.6', 'Suboptimal', 'Acceptable, minor delay'],
        ['2.5s', '60.7', 'Suboptimal', 'Noticeable but tolerable'],
        ['3.0s', '48.7', 'Disruptive', 'Cognitive load increases'],
        ['4.0s', '27.8', 'Disruptive', 'Significant misalignment'],
        ['5.0s', '13.5', 'Unacceptable (>5s)', 'Learner struggles to connect'],
        ['7.0s', '1.6', 'Unacceptable', 'Essentially unrelated in time'],
        ['10.0s', '0.0003', 'Unacceptable', 'Complete temporal disconnect'],
    ],
    col_widths=[0.8, 0.9, 1.3, 3.0]
)

add_heading('9.2 Zone Classification', 2)
add_table(
    ['Zone', 'Threshold', 'Cognitive Meaning'],
    [
        ['Optimal', '|delta_t| \u2264 1.0s', 'Visual and audio are perceived as simultaneous'],
        ['Suboptimal', '|delta_t| \u2264 3.0s', 'Small delay; learner can bridge the gap'],
        ['Disruptive', '|delta_t| \u2264 5.0s', 'Working memory strained; learning impaired'],
        ['Unacceptable', '|delta_t| > 5.0s', 'Split attention; learner cannot connect visual and audio'],
    ],
    col_widths=[1.2, 1.5, 3.7]
)

add_heading('9.3 Overall Score & Grade', 2)
add_formula('overall_score = \u03a3(S_temporal \u00d7 \u03b1) / \u03a3(\u03b1)')
add_para('(Alpha-weighted mean across all matched scenes)')

add_table(
    ['Score Range', 'Grade'],
    [
        ['\u2265 80', 'Excellent'],
        ['60 \u2013 79', 'Good'],
        ['40 \u2013 59', 'Acceptable'],
        ['20 \u2013 39', 'Poor'],
        ['< 20', 'Unacceptable'],
    ],
    col_widths=[2.0, 2.0]
)

add_heading('9.4 Expected Output', 2)
add_para('File: scores_per_scene.csv', bold=True)
add_para('Contains all alignment_events columns plus S_temporal and zone.')

add_para('\nFile: results.json (aggregates)', bold=True)
add_bullet('n_scenes, n_matched, n_no_match')
add_bullet('mean_S_temporal, median_S_temporal')
add_bullet('mean_delta_t, sd_delta_t, min_delta_t, max_delta_t')
add_bullet('Zone distribution: pct_Optimal, pct_Suboptimal, pct_Disruptive, pct_Unacceptable')
add_bullet('overall_score, overall_grade')
add_bullet('Keyword-level: n_keywords_total, n_grounded, grounding method distribution')

doc.add_page_break()

# ============================================================
# EXPECTED OUTPUTS SUMMARY
# ============================================================
add_heading('10. Complete Output File Inventory', 1)

add_para('For each video, the pipeline produces these files in the output directory:')

add_table(
    ['File', 'Stage', 'Rows (typical)', 'Key Information'],
    [
        ['transcript_words.csv', '1', '400\u2013800', 'Per-word timestamps'],
        ['transcript_segments.csv', '1', '30\u201380', 'Raw Whisper segments'],
        ['transcript_segments_improved.csv', '1', '25\u201360', 'pySBD sentences'],
        ['scenes.csv', '2', '10\u201325', 'Scene boundaries + OCR words'],
        ['ocr_per_frame.csv', '2', '400+', 'OCR text per frame'],
        ['dinov2_distances.csv', '2', '400+', 'Visual change signal'],
        ['scene_concepts.csv', '3', '10\u201325', 'Content labels + frame types'],
        ['segment_meta.csv', '4a', '25\u201360', 'Segment midpoints'],
        ['segment_embeddings.npy', '4a', '\u2014', '(N, 1024) embedding matrix'],
        ['alignment_events.csv', '4b', '10\u201325', 'Per-scene alignment (delta_t)'],
        ['keyword_alignment.csv', '5', '100\u2013400', 'Per-keyword grounding'],
        ['segment_keyword_scores.csv', '5', '25\u201360', 'Segment-level keyword stats'],
        ['pedagogical_importance.csv', '6', '25\u201360', 'Importance ratings 1\u20135'],
        ['scores_per_scene.csv', '7', '10\u201325', 'S_temporal + zones'],
        ['scores_weighted.csv', '7', '10\u201325', 'Importance-weighted scores'],
        ['results.json', '7', '\u2014', 'Overall aggregates + grade'],
        ['report_dashboard.html', '\u2014', '\u2014', 'Interactive HTML report'],
        ['timings.json', '\u2014', '\u2014', 'Per-stage timing'],
    ],
    col_widths=[2.5, 0.6, 1.0, 2.3]
)

doc.add_page_break()

# ============================================================
# KNOWN ISSUES
# ============================================================
add_heading('11. Known Issues in Current Pipeline', 1)

add_para('The following issues were identified by comparing V2.1 GPU results with V2.0 baseline:')

add_table(
    ['#', 'Issue', 'Severity', 'Impact'],
    [
        ['P1', 'delta_t uses t_start but all tracks match at t_keyframe', 'CRITICAL', 'Inflates delta_t by 5\u201340s for long scenes'],
        ['P2', 'Scene detection creates 51\u201356s monster scenes', 'CRITICAL', 'Only 13 scenes instead of 16\u201326'],
        ['P3', 'OCR captured at keyframe only, not scene start', 'HIGH', 'Concept doesn\'t represent scene start content'],
        ['P4', 'Gaussian scoring harsher than V2.0 at d=2\u20134s', 'HIGH', '21-point penalty vs V2.0 at d=3s'],
        ['P5', 'SigLIP classification all ~0.51 (random chance)', 'MEDIUM', 'No useful content/non-content filtering'],
        ['P6', 'TEMPORAL_SIGMA=15s >> SCORE_TAU=2.5s', 'MEDIUM', 'Alignment accepts far matches; scoring kills them'],
        ['P7', 'SigLIP grounding threshold 0.15 too lenient', 'MEDIUM', '86% of keywords grounded via weak SigLIP match'],
        ['P8', 'OCR text quality issues', 'LOW-MED', '"Iocal" instead of "local", garbled text'],
        ['P9', 'No temporal ordering constraint', 'LOW', 'Scenes can match out of temporal order'],
    ],
    col_widths=[0.4, 2.5, 0.8, 2.5]
)

add_heading('11.1 What V2.0 Does Right (That V2.1 Should Adopt)', 2)
add_bullet('SSIM-based scene detection produces 16 well-sized scenes (avg 8.9s)')
add_bullet('t_vis = t_start works because scenes are short (t_start \u2248 t_keyframe)')
add_bullet('Piecewise linear scoring is more forgiving: gives 70 at d=3s vs V2.1\'s 48.7')
add_bullet('S_final = S_raw \u00d7 (0.5 + 0.5 \u00d7 \u03b1) provides a 50% floor')
add_bullet('Tight \u00b15s search window (vs V2.1\'s \u03c3=15s)')

# Save
out_path = os.path.join(r"D:\PhD\Track2_codeClaude_19Mar26", "Pipeline_V2_1_Algorithm_Reference.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
